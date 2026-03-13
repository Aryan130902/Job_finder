"""
Resume Document Generator.

This module provides functionality to generate resume documents
in various formats (txt, docx) from ATS-optimized resume data.
"""

from typing import Optional
from core.optimizer.optimizer import ATSResume
import os


class DocumentGenerator:
    """
    Generator for creating resume documents in various formats.
    
    Supports generating resumes in text and DOCX formats
    from ATS-optimized resume data.
    """
    
    def __init__(self, ats_resume: ATSResume):
        """
        Initialize the document generator.
        
        Args:
            ats_resume: ATS-optimized resume object
        """
        self.ats_resume = ats_resume
    
    def to_text(self) -> str:
        """
        Generate text representation of the resume.
        
        Returns:
            Resume as formatted text string
        """
        lines = []
        
        lines.append(self.ats_resume.name.upper())
        lines.append("-" * 50)
        
        contact_info = []
        contact_info.append(f"Phone: {self.ats_resume.phone}")
        contact_info.append(f"Email: {self.ats_resume.email}")
        if self.ats_resume.linkedin:
            contact_info.append(f"LinkedIn: {self.ats_resume.linkedin}")
        
        lines.append(" | ".join(contact_info))
        lines.append("")
        
        if self.ats_resume.summary:
            lines.append("PROFESSIONAL SUMMARY")
            lines.append("-" * 50)
            lines.append(self.ats_resume.summary)
            lines.append("")
        
        if self.ats_resume.experience:
            lines.append("EXPERIENCE")
            lines.append("-" * 50)
            for exp in self.ats_resume.experience:
                lines.append(f"{exp['role']} at {exp['company']}")
                lines.append(f"{exp['duration']} | {exp['location']}")
                for bullet in exp['bullet_points']:
                    lines.append(f"  • {bullet}")
                lines.append("")
        
        if self.ats_resume.projects:
            lines.append("PROJECTS")
            lines.append("-" * 50)
            for proj in self.ats_resume.projects:
                lines.append(proj['name'])
                if proj['tech_stack']:
                    lines.append(f"Tech Stack: {proj['tech_stack']}")
                if proj['link']:
                    lines.append(f"Link: {proj['link']}")
                for bullet in proj['bullet_points']:
                    lines.append(f"  • {bullet}")
                lines.append("")
        
        if self.ats_resume.skills:
            lines.append("SKILLS")
            lines.append("-" * 50)
            for category, skills in self.ats_resume.skills.items():
                lines.append(f"{category}: {', '.join(skills)}")
            lines.append("")
        
        if self.ats_resume.education:
            lines.append("EDUCATION")
            lines.append("-" * 50)
            for edu in self.ats_resume.education:
                degree_info = edu['degree']
                if edu.get('cgpa'):
                    degree_info += f" | CGPA: {edu['cgpa']}"
                lines.append(f"{edu['institution']}")
                lines.append(f"{degree_info} | {edu['year']}")
                lines.append("")
        
        return "\n".join(lines)
    
    def to_docx(self, output_path: str) -> bool:
        """
        Generate DOCX document.
        
        Args:
            output_path: Path to save the DOCX file
            
        Returns:
            True if successful, False otherwise
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            print("python-docx not installed. Please install it with: pip install python-docx")
            return False
        
        doc = Document()
        
        title = doc.add_heading(self.ats_resume.name, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_parts = [
            self.ats_resume.phone,
            self.ats_resume.email,
        ]
        if self.ats_resume.linkedin:
            contact_parts.append(self.ats_resume.linkedin)
        contact_para.add_run(" | ".join(contact_parts))
        
        if self.ats_resume.summary:
            doc.add_heading("Professional Summary", level=1)
            doc.add_paragraph(self.ats_resume.summary)
        
        if self.ats_resume.experience:
            doc.add_heading("Experience", level=1)
            for exp in self.ats_resume.experience:
                heading = doc.add_heading(f"{exp['role']} at {exp['company']}", level=2)
                subhead = doc.add_paragraph(f"{exp['duration']} | {exp['location']}")
                subhead.runs[0].italic = True
                
                for bullet in exp['bullet_points']:
                    doc.add_paragraph(bullet, style='List Bullet')
        
        if self.ats_resume.projects:
            doc.add_heading("Projects", level=1)
            for proj in self.ats_resume.projects:
                heading = doc.add_heading(proj['name'], level=2)
                if proj.get('tech_stack'):
                    doc.add_paragraph(f"Tech Stack: {proj['tech_stack']}")
                if proj.get('link'):
                    doc.add_paragraph(f"Link: {proj['link']}")
                
                for bullet in proj['bullet_points']:
                    doc.add_paragraph(bullet, style='List Bullet')
        
        if self.ats_resume.skills:
            doc.add_heading("Skills", level=1)
            for category, skills in self.ats_resume.skills.items():
                doc.add_paragraph(f"{category}: {', '.join(skills)}")
        
        if self.ats_resume.education:
            doc.add_heading("Education", level=1)
            for edu in self.ats_resume.education:
                degree_info = edu['degree']
                if edu.get('cgpa'):
                    degree_info += f" | CGPA: {edu['cgpa']}"
                doc.add_paragraph(degree_info)
                doc.add_paragraph(f"{edu['institution']} | {edu['year']}")
        
        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        doc.save(output_path)
        return True
    
    def save(self, output_path: str, format: str = "txt") -> bool:
        """
        Save resume to file in specified format.
        
        Args:
            output_path: Path to save the file
            format: Output format - "txt" or "docx"
            
        Returns:
            True if successful, False otherwise
            
        Raises:
            ValueError: If format is not supported
        """
        if format == "txt":
            content = self.to_text()
            os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            return True
        elif format == "docx":
            return self.to_docx(output_path)
        else:
            raise ValueError(f"Unsupported format: {format}")


def generate_resume_document(
    ats_resume: ATSResume,
    output_path: str,
    format: str = "txt"
) -> bool:
    """
    Convenience function to generate a resume document.
    
    Args:
        ats_resume: ATS-optimized resume object
        output_path: Path to save the file
        format: Output format - "txt" or "docx"
        
    Returns:
        True if successful, False otherwise
    """
    generator = DocumentGenerator(ats_resume)
    return generator.save(output_path, format)
